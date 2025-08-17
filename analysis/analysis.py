import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from pathlib import Path
import numpy as np

class Analysis:
    def __init__(self, data_path='results/iterations/results.csv', output_dir='analysis/figures'):
        self.data_path = data_path
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(exist_ok=True)
        self.df = pd.read_csv(data_path)

        self.paradigm_colors = {
            'naive': '#FF6B6B',
            'advanced': '#4ECDC4',
            'agentic': '#45B7D1'
        }

        self.model_colors = {
            'gpt-4': '#2E86AB',
            'deepseek-chat': '#E07B91'
        }

        sns.set(style='whitegrid')
        print("Data loaded:", self.df.shape)

    def buildChartA(self):
        """Chart A: Question-Level Analysis - Performance Distribution by Complexity"""

        fig, ax = plt.subplots(1, 1, figsize=(12, 3))
        
        sns.boxplot(data=self.df, x='complexity', y='answer_correctness', hue='paradigm', ax=ax, palette=self.paradigm_colors)
        ax.set_ylabel('Answer Correctness Distribution', fontsize=12)
        ax.set_xlabel('Question Complexity', fontsize=12)
        ax.set_title('Performance Distribution by Question Complexity', fontweight='bold', fontsize=14)
        
        ax.tick_params(axis='x', rotation=0, labelsize=11)
        ax.tick_params(axis='y', labelsize=11)
        ax.legend(title='RAG Paradigm', bbox_to_anchor=(1.05, 1), loc='upper left')

        plt.tight_layout()
        plt.savefig(self.output_dir / 'chartA.png', format='png', bbox_inches='tight', dpi=300)
        plt.close()

    def buildChartB(self):
        """Chart B: Win Rate Analysis - Win Rate by Question Complexity"""

        df = self.df.copy()
        
        if 'question_id' not in df.columns:
            unique_questions = df['question'].unique()
            question_id_map = {q: i for i, q in enumerate(unique_questions)}
            df['question_id'] = df['question'].map(question_id_map)
        
        df['max_score'] = df.groupby('question_id')['semantic_similarity'].transform('max')
        df['is_winner'] = df['semantic_similarity'] == df['max_score']
    
        tie_counts = df[df['is_winner']].groupby('question_id').size()
        tied_questions = tie_counts[tie_counts > 1].index
        
        for q_id in tied_questions:
            tie_size = tie_counts[q_id]
            df.loc[(df['question_id'] == q_id) & df['is_winner'], 'win_weight'] = 1.0 / tie_size
        
        df.loc[df['is_winner'] & (~df['question_id'].isin(tied_questions)), 'win_weight'] = 1.0
        df['win_weight'] = df['win_weight'].fillna(0.0)
        
        fig, ax = plt.subplots(1, 1, figsize=(12, 3))
        
        win_analysis = []
        for complexity in df['complexity'].unique():
            complexity_data = df[df['complexity'] == complexity]
            total_questions = complexity_data['question_id'].nunique()
            
            for paradigm in df['paradigm'].unique():
                paradigm_complexity = complexity_data[complexity_data['paradigm'] == paradigm]
                wins = paradigm_complexity['win_weight'].sum()
                win_rate = (wins / total_questions) * 100 if total_questions > 0 else 0
                
                win_analysis.append({
                    'complexity': complexity,
                    'paradigm': paradigm,
                    'wins': wins,
                    'total_questions': total_questions,
                    'win_rate': win_rate
                })
        
        win_df = pd.DataFrame(win_analysis)
        
        sns.barplot(data=win_df, x='complexity', y='win_rate', hue='paradigm', palette=self.paradigm_colors, ax=ax)
        ax.set_ylabel('Win Rate (%)', fontsize=12)
        ax.set_xlabel('Question Complexity', fontsize=12)
        ax.set_title('Win Rate by Question Complexity', fontweight='bold', fontsize=14)

        max_win_rate = win_df['win_rate'].max()
        ax.set_ylim(0, min(100, max_win_rate * 1.1)) 
        ax.tick_params(axis='x', rotation=0, labelsize=10)
        ax.tick_params(axis='y', labelsize=11)
    
        plt.setp(ax.get_xticklabels(), ha='center')
        
        for container in ax.containers:
            ax.bar_label(container, fmt='%.1f%%', fontsize=9)
        
        ax.legend(title='RAG Paradigm', bbox_to_anchor=(1.05, 1), loc='upper left')
        
        plt.tight_layout()
        plt.savefig(self.output_dir / 'chartB.png', format='png', bbox_inches='tight', dpi=300)
        plt.close()

    def buildTableA(self):
        """Table A: Resource Efficiency Matrix - Cost, latency, tokens, and efficiency metrics by paradigm"""
        
        resource_metrics = ['cost_usd', 'latency', 'tokens_used']
        performance_metric = 'semantic_similarity' 
        available_resources = [metric for metric in resource_metrics if metric in self.df.columns]
        
        if performance_metric not in self.df.columns:
            print(f"Warning: {performance_metric} not found. Cannot calculate efficiency metrics.")
            return None
        
        efficiency_data = []
        
        for paradigm in self.df['paradigm'].unique():
            paradigm_data = self.df[self.df['paradigm'] == paradigm]
            
            row = {'Paradigm': paradigm.capitalize()}
            
            if 'cost_usd' in available_resources:
                row['Avg Cost ($)'] = round(paradigm_data['cost_usd'].mean(), 4)
            
            if 'latency' in available_resources:
                row['Avg Latency (s)'] = round(paradigm_data['latency'].mean(), 1)
            
            if 'tokens_used' in available_resources:
                row['Avg Tokens'] = round(paradigm_data['tokens_used'].mean(), 0)
            
            avg_performance = paradigm_data[performance_metric].mean()
            
            if 'cost_usd' in available_resources:
                avg_cost = paradigm_data['cost_usd'].mean()
                if avg_cost > 0:
                    perf_per_cost = avg_performance / avg_cost
                    row['Perf/Cost'] = round(perf_per_cost, 1)
                else:
                    row['Perf/Cost'] = None
            
            if 'tokens_used' in available_resources:
                avg_tokens = paradigm_data['tokens_used'].mean()
                if avg_tokens > 0:
                    perf_per_token = avg_performance / avg_tokens
                    row['Perf/Token'] = round(perf_per_token, 6)
                else:
                    row['Perf/Token'] = None
            
            if 'latency' in available_resources:
                avg_latency = paradigm_data['latency'].mean()
                if avg_latency > 0:
                    perf_per_second = avg_performance / avg_latency
                    row['Perf/Second'] = round(perf_per_second, 4)
                else:
                    row['Perf/Second'] = None
            
            row['Sample Size'] = len(paradigm_data)
            
            efficiency_data.append(row)
        
        efficiency_df = pd.DataFrame(efficiency_data)
        overall_row = {'Paradigm': 'Overall'}
        
        if 'cost_usd' in available_resources:
            overall_row['Avg Cost ($)'] = round(self.df['cost_usd'].mean(), 4)
        
        if 'latency' in available_resources:
            overall_row['Avg Latency (s)'] = round(self.df['latency'].mean(), 1)
        
        if 'tokens_used' in available_resources:
            overall_row['Avg Tokens'] = round(self.df['tokens_used'].mean(), 0)
        
        overall_performance = self.df[performance_metric].mean()
        
        if 'cost_usd' in available_resources:
            overall_cost = self.df['cost_usd'].mean()
            if overall_cost > 0:
                overall_row['Perf/Cost'] = round(overall_performance / overall_cost, 1)
        
        if 'tokens_used' in available_resources:
            overall_tokens = self.df['tokens_used'].mean()
            if overall_tokens > 0:
                overall_row['Perf/Token'] = round(overall_performance / overall_tokens, 6)
        
        if 'latency' in available_resources:
            overall_latency = self.df['latency'].mean()
            if overall_latency > 0:
                overall_row['Perf/Second'] = round(overall_performance / overall_latency, 4)
        
        overall_row['Sample Size'] = len(self.df)
        
        overall_df = pd.DataFrame([overall_row])
        efficiency_df = pd.concat([efficiency_df, overall_df], ignore_index=True)
        
        if 'Perf/Cost' in efficiency_df.columns:
            cost_ranking = efficiency_df[efficiency_df['Paradigm'] != 'Overall'].sort_values('Perf/Cost', ascending=False)
            print("\nCost Efficiency (Performance per Dollar):")
            for i, (_, row) in enumerate(cost_ranking.iterrows(), 1):
                if pd.notna(row['Perf/Cost']):
                    print(f"  {i}. {row['Paradigm']}: {row['Perf/Cost']}")
        
        if 'Perf/Token' in efficiency_df.columns:
            token_ranking = efficiency_df[efficiency_df['Paradigm'] != 'Overall'].sort_values('Perf/Token', ascending=False)
            print("\nToken Efficiency (Performance per Token):")
            for i, (_, row) in enumerate(token_ranking.iterrows(), 1):
                if pd.notna(row['Perf/Token']):
                    print(f"  {i}. {row['Paradigm']}: {row['Perf/Token']:.6f}")
        
        if 'Perf/Second' in efficiency_df.columns:
            speed_ranking = efficiency_df[efficiency_df['Paradigm'] != 'Overall'].sort_values('Perf/Second', ascending=False)
            print("\nSpeed Efficiency (Performance per Second):")
            for i, (_, row) in enumerate(speed_ranking.iterrows(), 1):
                if pd.notna(row['Perf/Second']):
                    print(f"  {i}. {row['Paradigm']}: {row['Perf/Second']:.4f}")
        
        csv_filename = self.output_dir / 'TableA_Resource_Efficiency.csv'
        efficiency_df.to_csv(csv_filename, index=False)
        word_filename = self.output_dir / 'TableA_Resource_Efficiency.docx'
        
        try:
            import docx
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            doc.add_heading('Table A: Resource Efficiency Matrix', level=1)
            
            table = doc.add_table(rows=1, cols=len(efficiency_df.columns))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(efficiency_df.columns):
                hdr_cells[i].text = str(column_name)
            
            for _, row in efficiency_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
            
            doc.save(word_filename)
            print(f"\nWord document exported to: {word_filename}")
            
        except ImportError as e:
            print(f"\npython-docx not available: {e}")
            word_filename = None
        except Exception as e:
            print(f"\nError creating Word document: {e}")
            word_filename = None
        
        print(f"CSV exported to: {csv_filename}")
        
        return {
            'table': efficiency_df,
            'csv_filename': csv_filename,
            'word_filename': word_filename,
            'available_resources': available_resources
        }

    
        """Table D: Performance by Model - Clear comparison of how each paradigm performs with different models"""
        
        # Check if we have the required columns
        if 'model' not in self.df.columns or 'semantic_similarity' not in self.df.columns:
            print("Error: 'model' and 'semantic_similarity' columns required for model analysis")
            return None
        
        available_models = self.df['model'].unique()
        print(f"Available models: {available_models}")
        
        # Create model name mapping for cleaner display
        model_names = {}
        for model in available_models:
            if 'gpt-4' in model.lower():
                model_names[model] = 'GPT-4'
            elif 'deepseek' in model.lower():
                model_names[model] = 'DeepSeek'
            else:
                model_names[model] = model.replace('_', ' ').title()
        
        # Define key metrics to show
        key_metrics = {
            'semantic_similarity': 'Semantic Similarity',
            'faithfulness': 'Faithfulness',
            'answer_correctness': 'Answer Correctness',
            'cost_usd': 'Avg Cost ($)',
            'latency': 'Avg Latency (s)'
        }
        
        # Check which metrics are available
        available_metrics = {k: v for k, v in key_metrics.items() if k in self.df.columns}
        
        # Calculate performance for each paradigm-model combination
        performance_data = []
        
        for paradigm in self.df['paradigm'].unique():
            for model in available_models:
                subset = self.df[(self.df['paradigm'] == paradigm) & (self.df['model'] == model)]
                
                if len(subset) > 0:
                    row = {
                        'Paradigm': paradigm.capitalize(),
                        'Model': model_names[model],
                        'Questions': len(subset)
                    }
                    
                    # Add performance metrics
                    for metric_key, metric_label in available_metrics.items():
                        if metric_key in subset.columns:
                            mean_val = subset[metric_key].mean()
                            
                            # Format based on metric type
                            if metric_key == 'cost_usd':
                                row[metric_label] = round(mean_val, 4)
                            elif metric_key == 'latency':
                                row[metric_label] = round(mean_val, 1)
                            else:
                                row[metric_label] = round(mean_val, 3)
                        else:
                            row[metric_label] = None
                    
                    performance_data.append(row)
        
        # Create DataFrame
        performance_df = pd.DataFrame(performance_data)
        
        # Print table to console
        print("\n" + "="*120)
        print("TABLE D: PERFORMANCE BY MODEL")
        print("="*120)
        print()
        print(performance_df.to_string(index=False))
        print("="*120)
        
        # Analysis insights
        print(f"\n{'-'*80}")
        print("MODEL PERFORMANCE INSIGHTS")
        print(f"{'-'*80}")
        
        # Best performing paradigm per model (using semantic similarity)
        if 'Semantic Similarity' in performance_df.columns:
            print(f"\nBest Paradigm per Model (by Semantic Similarity):")
            for model_name in model_names.values():
                model_data = performance_df[performance_df['Model'] == model_name]
                if len(model_data) > 0 and model_data['Semantic Similarity'].notna().any():
                    best_row = model_data.loc[model_data['Semantic Similarity'].idxmax()]
                    print(f"  {model_name}: {best_row['Paradigm']} ({best_row['Semantic Similarity']:.3f})")
        
        # Best model per paradigm
        print(f"\nBest Model per Paradigm (by Semantic Similarity):")
        for paradigm in self.df['paradigm'].unique():
            paradigm_cap = paradigm.capitalize()
            paradigm_data = performance_df[performance_df['Paradigm'] == paradigm_cap]
            if len(paradigm_data) > 0 and 'Semantic Similarity' in paradigm_data.columns:
                if paradigm_data['Semantic Similarity'].notna().any():
                    best_row = paradigm_data.loc[paradigm_data['Semantic Similarity'].idxmax()]
                    print(f"  {paradigm_cap}: {best_row['Model']} ({best_row['Semantic Similarity']:.3f})")
        
        # Performance differences between models
        if len(available_models) >= 2 and 'Semantic Similarity' in performance_df.columns:
            print(f"\nPerformance Differences Between Models:")
            for paradigm in self.df['paradigm'].unique():
                paradigm_cap = paradigm.capitalize()
                paradigm_data = performance_df[performance_df['Paradigm'] == paradigm_cap]
                
                if len(paradigm_data) >= 2:
                    paradigm_data_clean = paradigm_data[paradigm_data['Semantic Similarity'].notna()]
                    if len(paradigm_data_clean) >= 2:
                        max_perf = paradigm_data_clean['Semantic Similarity'].max()
                        min_perf = paradigm_data_clean['Semantic Similarity'].min()
                        difference = max_perf - min_perf
                        
                        best_model = paradigm_data_clean.loc[paradigm_data_clean['Semantic Similarity'].idxmax(), 'Model']
                        worst_model = paradigm_data_clean.loc[paradigm_data_clean['Semantic Similarity'].idxmin(), 'Model']
                        
                        print(f"  {paradigm_cap}: {difference:.3f} gap ({best_model}: {max_perf:.3f} vs {worst_model}: {min_perf:.3f})")
        
        # Cost and speed analysis if available
        if 'Avg Cost ($)' in performance_df.columns:
            print(f"\nCost Analysis:")
            print(f"Most Expensive: ", end="")
            most_expensive = performance_df.loc[performance_df['Avg Cost ($)'].idxmax()]
            print(f"{most_expensive['Paradigm']} with {most_expensive['Model']} (${most_expensive['Avg Cost ($)']:.4f})")
            
            print(f"Most Affordable: ", end="")
            most_affordable = performance_df.loc[performance_df['Avg Cost ($)'].idxmin()]
            print(f"{most_affordable['Paradigm']} with {most_affordable['Model']} (${most_affordable['Avg Cost ($)']:.4f})")
        
        if 'Avg Latency (s)' in performance_df.columns:
            print(f"\nSpeed Analysis:")
            print(f"Fastest: ", end="")
            fastest = performance_df.loc[performance_df['Avg Latency (s)'].idxmin()]
            print(f"{fastest['Paradigm']} with {fastest['Model']} ({fastest['Avg Latency (s)']:.1f}s)")
            
            print(f"Slowest: ", end="")
            slowest = performance_df.loc[performance_df['Avg Latency (s)'].idxmax()]
            print(f"{slowest['Paradigm']} with {slowest['Model']} ({slowest['Avg Latency (s)']:.1f}s)")
        
        # Export to CSV
        csv_filename = self.output_dir / 'TableD_Performance_by_Model.csv'
        performance_df.to_csv(csv_filename, index=False)
        
        # Try to create Word document
        word_filename = self.output_dir / 'TableD_Performance_by_Model.docx'
        
        try:
            import docx
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            # Create Word document
            doc = Document()
            doc.add_heading('Table D: Performance by Model', level=1)
            
            # Create table
            table = doc.add_table(rows=1, cols=len(performance_df.columns))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            # Add header row
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(performance_df.columns):
                hdr_cells[i].text = str(column_name)
            
            # Add data rows
            for _, row in performance_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
            
            # Save Word document
            doc.save(word_filename)
            print(f"\nWord document exported to: {word_filename}")
            
        except ImportError as e:
            print(f"\npython-docx not available: {e}")
            word_filename = None
        except Exception as e:
            print(f"\nError creating Word document: {e}")
            word_filename = None
        
        print(f"CSV exported to: {csv_filename}")
        
        return {
            'table': performance_df,
            'csv_filename': csv_filename,
            'word_filename': word_filename,
            'available_models': available_models,
            'model_names': model_names
        }

    def buildTableB(self):
        """Table B: Resource Efficiency Matrix - Clear comparison of how each paradigm performs with different models"""

        if 'model' not in self.df.columns or 'semantic_similarity' not in self.df.columns:
            print("Error: 'model' and 'semantic_similarity' columns required for model analysis")
            return None
        
        available_models = self.df['model'].unique()
        print(f"Available models: {available_models}")
        
        model_names = {}
        for model in available_models:
            if 'gpt-4' in model.lower():
                model_names[model] = 'GPT-4'
            elif 'deepseek' in model.lower():
                model_names[model] = 'DeepSeek'
            else:
                model_names[model] = model.replace('_', ' ').title()
        
        key_metrics = {
            'semantic_similarity': 'Semantic Similarity',
            'faithfulness': 'Faithfulness',
            'answer_correctness': 'Answer Correctness',
            'context_precision': 'Context Precision',
            'tokens_used': 'Total Tokens',
            'cost_usd': 'Avg Cost ($)',
            'latency': 'Avg Latency (s)'
        }
        
        available_metrics = {k: v for k, v in key_metrics.items() if k in self.df.columns}
        performance_data = []
        
        for paradigm in self.df['paradigm'].unique():
            for model in available_models:
                subset = self.df[(self.df['paradigm'] == paradigm) & (self.df['model'] == model)]
                
                if len(subset) > 0:
                    row = {
                        'Paradigm': paradigm.capitalize(),
                        'Model': model_names[model],
                        'Questions': len(subset)
                    }
                    
                    for metric_key, metric_label in available_metrics.items():
                        if metric_key in subset.columns:
                            mean_val = subset[metric_key].mean()
                            
                            if metric_key == 'cost_usd':
                                row[metric_label] = round(mean_val, 4)
                            elif metric_key == 'latency':
                                row[metric_label] = round(mean_val, 1)
                            elif metric_key == 'tokens_used':
                                row[metric_label] = round(mean_val, 0) 
                            else:
                                row[metric_label] = round(mean_val, 3)
                        else:
                            row[metric_label] = None
                    
                    performance_data.append(row)
        
        performance_df = pd.DataFrame(performance_data)
        
        if 'Semantic Similarity' in performance_df.columns:
            print(f"\nBest Paradigm per Model (by Semantic Similarity):")
            for model_name in model_names.values():
                model_data = performance_df[performance_df['Model'] == model_name]
                if len(model_data) > 0 and model_data['Semantic Similarity'].notna().any():
                    best_row = model_data.loc[model_data['Semantic Similarity'].idxmax()]
                    print(f"  {model_name}: {best_row['Paradigm']} ({best_row['Semantic Similarity']:.3f})")
        
        print(f"\nBest Model per Paradigm (by Semantic Similarity):")
        for paradigm in self.df['paradigm'].unique():
            paradigm_cap = paradigm.capitalize()
            paradigm_data = performance_df[performance_df['Paradigm'] == paradigm_cap]
            if len(paradigm_data) > 0 and 'Semantic Similarity' in paradigm_data.columns:
                if paradigm_data['Semantic Similarity'].notna().any():
                    best_row = paradigm_data.loc[paradigm_data['Semantic Similarity'].idxmax()]
                    print(f"  {paradigm_cap}: {best_row['Model']} ({best_row['Semantic Similarity']:.3f})")
        
        if len(available_models) >= 2 and 'Semantic Similarity' in performance_df.columns:
            print(f"\nPerformance Differences Between Models:")
            for paradigm in self.df['paradigm'].unique():
                paradigm_cap = paradigm.capitalize()
                paradigm_data = performance_df[performance_df['Paradigm'] == paradigm_cap]
                
                if len(paradigm_data) >= 2:
                    paradigm_data_clean = paradigm_data[paradigm_data['Semantic Similarity'].notna()]
                    if len(paradigm_data_clean) >= 2:
                        max_perf = paradigm_data_clean['Semantic Similarity'].max()
                        min_perf = paradigm_data_clean['Semantic Similarity'].min()
                        difference = max_perf - min_perf
                        
                        best_model = paradigm_data_clean.loc[paradigm_data_clean['Semantic Similarity'].idxmax(), 'Model']
                        worst_model = paradigm_data_clean.loc[paradigm_data_clean['Semantic Similarity'].idxmin(), 'Model']
                        
                        print(f"  {paradigm_cap}: {difference:.3f} gap ({best_model}: {max_perf:.3f} vs {worst_model}: {min_perf:.3f})")
        
        if 'Avg Cost ($)' in performance_df.columns:
            print(f"\nCost Analysis:")
            print(f"Most Expensive: ", end="")
            most_expensive = performance_df.loc[performance_df['Avg Cost ($)'].idxmax()]
            print(f"{most_expensive['Paradigm']} with {most_expensive['Model']} (${most_expensive['Avg Cost ($)']:.4f})")
            
            print(f"Most Affordable: ", end="")
            most_affordable = performance_df.loc[performance_df['Avg Cost ($)'].idxmin()]
            print(f"{most_affordable['Paradigm']} with {most_affordable['Model']} (${most_affordable['Avg Cost ($)']:.4f})")
        
        if 'Avg Latency (s)' in performance_df.columns:
            print(f"\nSpeed Analysis:")
            print(f"Fastest: ", end="")
            fastest = performance_df.loc[performance_df['Avg Latency (s)'].idxmin()]
            print(f"{fastest['Paradigm']} with {fastest['Model']} ({fastest['Avg Latency (s)']:.1f}s)")
            
            print(f"Slowest: ", end="")
            slowest = performance_df.loc[performance_df['Avg Latency (s)'].idxmax()]
            print(f"{slowest['Paradigm']} with {slowest['Model']} ({slowest['Avg Latency (s)']:.1f}s)")
        
        csv_filename = self.output_dir / 'TableB_Performance_by_Model.csv'
        performance_df.to_csv(csv_filename, index=False)
        word_filename = self.output_dir / 'TableB_Performance_by_Model.docx'
        
        try:
            import docx
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            doc.add_heading('Table D: Performance by Model', level=1)
            
            table = doc.add_table(rows=1, cols=len(performance_df.columns))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(performance_df.columns):
                hdr_cells[i].text = str(column_name)
            
            for _, row in performance_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
            
            doc.save(word_filename)
            print(f"\nWord document exported to: {word_filename}")
            
        except ImportError as e:
            print(f"\npython-docx not available: {e}")
            word_filename = None
        except Exception as e:
            print(f"\nError creating Word document: {e}")
            word_filename = None
        
        print(f"CSV exported to: {csv_filename}")
        
        return {
            'table': performance_df,
            'csv_filename': csv_filename,
            'word_filename': word_filename,
            'available_models': available_models,
            'model_names': model_names
        }
    
    def buildTableC(self):
        """Table C: Overall Paradigm Rankings - Rankings across different metrics with performance gaps"""

        ranking_metrics = {
            'semantic_similarity': ('Semantic Similarity', 'higher'),
            'faithfulness': ('Faithfulness', 'higher'),
            'context_precision': ('Context Precision', 'higher'),
            'answer_correctness': ('Answer Correctness', 'higher'),
            'cost_usd': ('Cost Efficiency', 'lower'),  
            'latency': ('Speed', 'lower'), 
            'tokens_used': ('Token Efficiency', 'lower') 
        }
        
        consistency_metrics = ['semantic_similarity', 'faithfulness', 'context_precision']
        available_metrics = {k: v for k, v in ranking_metrics.items() if k in self.df.columns}
        
        if not available_metrics:
            print("Error: No suitable metrics found for ranking analysis")
            return None
        
        ranking_data = []
        
        for metric_key, (metric_label, direction) in available_metrics.items():
            paradigm_means = self.df.groupby('paradigm')[metric_key].mean().sort_values(ascending=(direction=='lower'))
            row = {'Metric': metric_label}
            
            for i, (paradigm, mean_value) in enumerate(paradigm_means.items()):
                position = ['1st Place', '2nd Place', '3rd Place'][i] if i < 3 else f'{i+1}th Place'
                row[position] = paradigm.capitalize()
            
            if len(paradigm_means) > 1:
                if direction == 'higher':
                    gap = paradigm_means.iloc[-1] - paradigm_means.iloc[0] 
                    row['Gap (1st-3rd)'] = round(gap, 3)
                else:
                    gap = paradigm_means.iloc[0] - paradigm_means.iloc[-1]  
                    row['Gap (1st-3rd)'] = round(gap, 3)
            else:
                row['Gap (1st-3rd)'] = 0
            
            ranking_data.append(row)
        
        for metric in consistency_metrics:
            if metric in self.df.columns:
                consistency_data = []
                for paradigm in self.df['paradigm'].unique():
                    paradigm_data = self.df[self.df['paradigm'] == paradigm]
                    mean_val = paradigm_data[metric].mean()
                    std_val = paradigm_data[metric].std()
                    cv = std_val / mean_val if mean_val > 0 else 0
                    consistency_data.append((paradigm, cv))
                
                consistency_data.sort(key=lambda x: x[1])
                
                metric_name = ranking_metrics.get(metric, (metric, 'higher'))[0]
                row = {'Metric': f'{metric_name} Consistency'}
                
                for i, (paradigm, cv) in enumerate(consistency_data):
                    position = ['1st Place', '2nd Place', '3rd Place'][i] if i < 3 else f'{i+1}th Place'
                    row[position] = paradigm.capitalize()
                
                if len(consistency_data) > 1:
                    cv_gap = consistency_data[-1][1] - consistency_data[0][1]
                    row['Gap (1st-3rd)'] = round(cv_gap, 3)
                else:
                    row['Gap (1st-3rd)'] = 0
                
                ranking_data.append(row)
        
        ranking_df = pd.DataFrame(ranking_data)
        paradigm_scores = {}

        for paradigm in self.df['paradigm'].unique():
            paradigm_cap = paradigm.capitalize()
            scores = {'1st': 0, '2nd': 0, '3rd': 0, 'total_score': 0}
            
            for _, row in ranking_df.iterrows():
                if row.get('1st Place') == paradigm_cap:
                    scores['1st'] += 1
                    scores['total_score'] += 3  
                elif row.get('2nd Place') == paradigm_cap:
                    scores['2nd'] += 1
                    scores['total_score'] += 2 
                elif row.get('3rd Place') == paradigm_cap:
                    scores['3rd'] += 1
                    scores['total_score'] += 1 
            
            paradigm_scores[paradigm_cap] = scores
        
        sorted_paradigms = sorted(paradigm_scores.items(), key=lambda x: x[1]['total_score'], reverse=True)
        
        print("Overall Rankings (3 pts for 1st, 2 pts for 2nd, 1 pt for 3rd):")
        for i, (paradigm, scores) in enumerate(sorted_paradigms, 1):
            print(f"  {i}. {paradigm}: {scores['total_score']} points ({scores['1st']} × 1st, {scores['2nd']} × 2nd, {scores['3rd']} × 3rd)")
        
        for paradigm_cap, scores in paradigm_scores.items():
            print(f"\n{paradigm_cap}:")
    
            strengths = []
            weaknesses = []
            
            for _, row in ranking_df.iterrows():
                metric = row['Metric']
                if row.get('1st Place') == paradigm_cap:
                    strengths.append(metric)
                elif row.get('3rd Place') == paradigm_cap:
                    weaknesses.append(metric)
            
            if strengths:
                print(f"  Strengths: {', '.join(strengths)}")
            if weaknesses:
                print(f"  Weaknesses: {', '.join(weaknesses)}")
            
            if not strengths and not weaknesses:
                print(f"  Consistent middle performer across metrics")
        
        total_metrics = len(ranking_df)
        
        for paradigm_cap, scores in paradigm_scores.items():
            dominance = (scores['1st'] / total_metrics) * 100 if total_metrics > 0 else 0
            print(f"{paradigm_cap}: Wins {scores['1st']}/{total_metrics} metrics ({dominance:.1f}% dominance)")
        
        csv_filename = self.output_dir / 'TableC_Overall_Rankings.csv'
        ranking_df.to_csv(csv_filename, index=False)
        summary_data = []

        for paradigm_cap, scores in sorted_paradigms:
            summary_data.append({
                'Paradigm': paradigm_cap,
                '1st Place Finishes': scores['1st'],
                '2nd Place Finishes': scores['2nd'], 
                '3rd Place Finishes': scores['3rd'],
                'Total Score': scores['total_score'],
                'Dominance (%)': round((scores['1st'] / total_metrics) * 100, 1) if total_metrics > 0 else 0
            })
        
        summary_df = pd.DataFrame(summary_data)
        summary_csv_filename = self.output_dir / 'TableC_Ranking_Summary.csv'
        summary_df.to_csv(summary_csv_filename, index=False)
        word_filename = self.output_dir / 'TableC_Overall_Rankings.docx'
        
        try:
            import docx
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            doc.add_heading('Table F: Overall Paradigm Rankings', level=1)
            
            table = doc.add_table(rows=1, cols=len(ranking_df.columns))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(ranking_df.columns):
                hdr_cells[i].text = str(column_name)
            
            for _, row in ranking_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
            
            doc.add_heading('Ranking Summary', level=2)
            
            summary_table = doc.add_table(rows=1, cols=len(summary_df.columns))
            summary_table.style = 'Light Grid Accent 1'
            summary_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = summary_table.rows[0].cells
            for i, column_name in enumerate(summary_df.columns):
                hdr_cells[i].text = str(column_name)
            
            for _, row in summary_df.iterrows():
                row_cells = summary_table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
            
            doc.save(word_filename)
            print(f"\nWord document exported to: {word_filename}")
            
        except ImportError as e:
            print(f"\npython-docx not available: {e}")
            word_filename = None
        except Exception as e:
            print(f"\nError creating Word document: {e}")
            word_filename = None
        
        print(f"Main rankings CSV exported to: {csv_filename}")
        print(f"Ranking summary CSV exported to: {summary_csv_filename}")
        
        return {
            'rankings_table': ranking_df,
            'summary_table': summary_df,
            'paradigm_scores': paradigm_scores,
            'csv_filename': csv_filename,
            'summary_csv_filename': summary_csv_filename,
            'word_filename': word_filename
        }

    def buildTableD(self):
        """Table D: Performance by Model - Performance metrics for each paradigm across different models"""
            
        if 'model' not in self.df.columns or 'semantic_similarity' not in self.df.columns:
            print("Error: 'model' and 'semantic_similarity' columns required for model analysis")
            return None
        
        available_models = self.df['model'].unique()
        print(f"Available models: {available_models}")
        performance_data = []
        
        for paradigm in self.df['paradigm'].unique():
            row = {'Paradigm': paradigm.capitalize()}
            
            for model in available_models:
                subset = self.df[(self.df['paradigm'] == paradigm) & (self.df['model'] == model)]
                
                if len(subset) > 0:
                    mean_perf = subset['semantic_similarity'].mean()
                    std_perf = subset['semantic_similarity'].std()
                    row[f'{model}'] = round(mean_perf, 3)
                    row[f'{model} (SD)'] = round(std_perf, 3)
                    row[f'{model} (n)'] = len(subset)
                else:
                    row[f'{model}'] = None
                    row[f'{model} (SD)'] = None
                    row[f'{model} (n)'] = 0
            
            if len(available_models) >= 2:
                paradigm_data = self.df[self.df['paradigm'] == paradigm]
                
                if 'question' in paradigm_data.columns:
                    pivot_data = paradigm_data.pivot_table(
                        index=['question'], 
                        columns='model', 
                        values='semantic_similarity'
                    ).dropna()
                    
                    if len(pivot_data) > 1 and len(pivot_data.columns) >= 2:
                        correlation = pivot_data.iloc[:, 0].corr(pivot_data.iloc[:, 1])
                        row['Model Agreement (r)'] = round(correlation, 3)
                    else:
                        row['Model Agreement (r)'] = None
                else:
                    row['Model Agreement (r)'] = None
            
            model_means = []
            for model in available_models:
                subset = self.df[(self.df['paradigm'] == paradigm) & (self.df['model'] == model)]
                if len(subset) > 0:
                    model_means.append(subset['semantic_similarity'].mean())
            
            if len(model_means) > 1:
                cv_across_models = np.std(model_means) / np.mean(model_means) if np.mean(model_means) > 0 else 0
                row['CV Across Models'] = round(cv_across_models, 3)
            else:
                row['CV Across Models'] = None
            
            performance_data.append(row)
        
        performance_df = pd.DataFrame(performance_data)
        
        if 'CV Across Models' in performance_df.columns:
            cv_data = performance_df[performance_df['CV Across Models'].notna()]
            if len(cv_data) > 0:
                most_consistent = cv_data.loc[cv_data['CV Across Models'].idxmin()]
                print(f"Most Consistent Paradigm: {most_consistent['Paradigm']} (CV = {most_consistent['CV Across Models']})")
                
                consistency_ranking = cv_data.sort_values('CV Across Models')[['Paradigm', 'CV Across Models']]
                print("\nConsistency Ranking (Lower CV = More Consistent):")
                for i, (_, row) in enumerate(consistency_ranking.iterrows(), 1):
                    print(f"  {i}. {row['Paradigm']}: {row['CV Across Models']}")
        
        if 'Model Agreement (r)' in performance_df.columns:
            agreement_data = performance_df[performance_df['Model Agreement (r)'].notna()]
            if len(agreement_data) > 0:
                highest_agreement = agreement_data.loc[agreement_data['Model Agreement (r)'].idxmax()]
                print(f"\nHighest Model Agreement: {highest_agreement['Paradigm']} (r = {highest_agreement['Model Agreement (r)']})")

                agreement_ranking = agreement_data.sort_values('Model Agreement (r)', ascending=False)[['Paradigm', 'Model Agreement (r)']]
                print("\nModel Agreement Ranking:")
                for i, (_, row) in enumerate(agreement_ranking.iterrows(), 1):
                    print(f"  {i}. {row['Paradigm']}: {row['Model Agreement (r)']}")
        
        print(f"\nBest Paradigm per Model:")
        for model in available_models:
            model_col = f'{model}'
            if model_col in performance_df.columns:
                model_data = performance_df[performance_df[model_col].notna()]
                if len(model_data) > 0:
                    best_paradigm = model_data.loc[model_data[model_col].idxmax()]
                    print(f"  {model}: {best_paradigm['Paradigm']} ({best_paradigm[model_col]})")
        
        csv_filename = self.output_dir / 'TableD_Performance_by_Model.csv'
        performance_df.to_csv(csv_filename, index=False)
        word_filename = self.output_dir / 'TableD_Performance_by_Model.docx'
        
        try:
            import docx
            from docx import Document
            from docx.enum.table import WD_TABLE_ALIGNMENT
            
            doc = Document()
            doc.add_heading('Table D: Performance by Model', level=1)
            
            table = doc.add_table(rows=1, cols=len(performance_df.columns))
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            hdr_cells = table.rows[0].cells
            for i, column_name in enumerate(performance_df.columns):
                hdr_cells[i].text = str(column_name)
            
            for _, row in performance_df.iterrows():
                row_cells = table.add_row().cells
                for i, value in enumerate(row):
                    row_cells[i].text = str(value) if value is not None else ""
            
            doc.save(word_filename)
            print(f"\nWord document exported to: {word_filename}")
            
        except ImportError as e:
            print(f"\npython-docx not available: {e}")
            word_filename = None
        except Exception as e:
            print(f"\nError creating Word document: {e}")
            word_filename = None
        
        print(f"CSV exported to: {csv_filename}")
        
        return {
            'table': performance_df,
            'csv_filename': csv_filename,
            'word_filename': word_filename,
            'available_models': available_models
        }

if __name__ == "__main__":
    analysis = Analysis()
    analysis.buildChartA()
    analysis.buildChartB()
    analysis.buildTableA()
    analysis.buildTableB()
    analysis.buildTableC()
    analysis.buildTableD()